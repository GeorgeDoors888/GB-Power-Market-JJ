from docx import Document

# Create a new Word document
doc = Document()

# Add a title
doc.add_heading("Report: Analysis of BSUoS Forecast and NESO Billing System", level=1)

# Add sections
# Section 1
section1 = doc.add_heading("1. BSUoS Forecast Analysis", level=2)
doc.add_paragraph("Key Metrics")
doc.add_paragraph("- Total Costs vs. Demand: Visualized trends over time.")
doc.add_paragraph("- Profit/Loss Trends: Highlighted monthly fluctuations.")

doc.add_paragraph("Visualizations")
doc.add_paragraph("- Line Plot: Total Costs vs. Demand.")
doc.add_paragraph("- Bar Plot: Profit/Loss trends.")

# Section 2
section2 = doc.add_heading("2. NESO Billing System Analysis", level=2)
doc.add_paragraph("Report Summary")
doc.add_paragraph("- Title: Detailed Report on BSUoS Billing System under NESO")
doc.add_paragraph("- Generated On: 2025-09-15")

doc.add_paragraph("Key Sections")
doc.add_paragraph("1. Overview of BSUoS Billing System")
doc.add_paragraph(
    "The Balancing Services Use of System (BSUoS) charges recover costs incurred by NESO in balancing the electricity system."
)

doc.add_paragraph("2. Billing Methodology and Systems")
doc.add_paragraph(
    "Since April 2023, BSUoS transitioned to a fixed-rate billing system for Final Demand users."
)

doc.add_paragraph("3. Participants and Settlement Routes")
doc.add_paragraph(
    "BSUoS is payable by suppliers and transmission-connected generators."
)

doc.add_paragraph("4. Timing of BSUoS Billing")
doc.add_paragraph(
    "Invoices are produced daily with reconciliations following a set schedule."
)

doc.add_paragraph("Insights")
doc.add_paragraph("- Transition to fixed-rate billing has streamlined processes.")
doc.add_paragraph("- Timing of reconciliations ensures accuracy in billing.")

# Save the document
doc.save("Updated_Report.docx")

print("Word document created successfully.")

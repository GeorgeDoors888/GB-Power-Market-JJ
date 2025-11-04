from __future__ import annotations
import io
from docx import Document

def extract_docx_text(blob: bytes) -> str:
    f = io.BytesIO(blob)
    doc = Document(f)
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)
